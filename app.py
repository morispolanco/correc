import streamlit as st
import requests
import io
from docx import Document
import re

st.title("Correcciones mínimas de texto en DOCX")

# Subir archivo DOCX
uploaded_file = st.file_uploader("Sube un archivo DOCX", type="docx")

if uploaded_file is not None:
    # Leer el documento DOCX
    document = Document(uploaded_file)
    
    # Extraer el texto del documento
    full_text = []
    for para in document.paragraphs:
        full_text.append(para.text)
    text = '\n'.join(full_text)
    
    # Función para procesar el texto sin afectar las citas textuales
    def process_text(text):
        # Encontrar citas textuales
        pattern = r'\".*?\"'
        quoted_texts = re.findall(pattern, text)
        
        # Reemplazar citas por un marcador
        temp_text = re.sub(pattern, '__CITA__', text)
        
        # Llamar a la API de Together para corregir el texto
        corrected_text = correct_text(temp_text)
        
        # Restaurar las citas textuales
        for quote in quoted_texts:
            corrected_text = corrected_text.replace('__CITA__', quote, 1)
        
        return corrected_text
    
    # Función para llamar a la API de Together
    def correct_text(input_text):
        api_url = "https://api.together.xyz/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {st.secrets['together_api_key']}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "meta-llama/Meta-Llama-3.1-405B-Instruct-Turbo",
            "messages": [
                {"role": "system", "content": "Eres un corrector ortográfico y gramatical. Haz correcciones mínimas al texto proporcionado, sin cambiar el significado y sin tocar o cambiar las citas textuales encerradas entre comillas."},
                {"role": "user", "content": input_text}
            ],
            "max_tokens": 2512,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["<|eot_id|>"],
            "stream": False
        }
        response = requests.post(api_url, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            corrected_text = result['choices'][0]['message']['content']
            return corrected_text
        else:
            st.error(f"Error en la API: {response.status_code}")
            st.stop()
    
    # Procesar el texto y obtener el texto corregido
    corrected_text = process_text(text)
    
    # Crear un nuevo documento con el texto corregido
    corrected_document = Document()
    for line in corrected_text.split('\n'):
        corrected_document.add_paragraph(line)
    
    # Guardar el documento corregido en un objeto BytesIO
    corrected_io = io.BytesIO()
    corrected_document.save(corrected_io)
    corrected_io.seek(0)
    
    # Botón para descargar el documento corregido
    st.download_button(
        label="Descargar documento corregido",
        data=corrected_io,
        file_name="documento_corregido.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
